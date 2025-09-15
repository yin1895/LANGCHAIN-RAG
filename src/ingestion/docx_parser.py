import hashlib
import logging
import re
from contextlib import contextmanager

# PDF 支持
from io import StringIO
from pathlib import Path
from typing import Dict, Iterable, List

from docx import Document

try:  # pdfminer 可能未安装
    from pdfminer.high_level import extract_text_to_fp

    HAVE_PDF = True
except Exception:  # pragma: no cover
    HAVE_PDF = False

# 轻量级 OMML 解析: 避免使用 xpath(namespaces=...) 以适配 python-docx BaseOxmlElement
# 做法：使用底层 XML 字符串 regex 捕获常见结构（目前仅分数，可扩展）

FRACTION_PATTERN = re.compile(r"<m:f>(.*?)</m:f>", re.DOTALL)
NUM_PATTERN = re.compile(r"<m:num>(.*?)</m:num>", re.DOTALL)
DEN_PATTERN = re.compile(r"<m:den>(.*?)</m:den>", re.DOTALL)
SUP_PATTERN = re.compile(r"<m:sSup>(.*?)</m:sSup>", re.DOTALL)
SUB_PATTERN = re.compile(r"<m:sSub>(.*?)</m:sSub>", re.DOTALL)
SUBSUP_PATTERN = re.compile(r"<m:sSubSup>(.*?)</m:sSubSup>", re.DOTALL)
RAD_PATTERN = re.compile(r"<m:rad>(.*?)</m:rad>", re.DOTALL)
NARY_PATTERN = re.compile(r"<m:nary>(.*?)</m:nary>", re.DOTALL)
TEXT_PATTERN = re.compile(r"<w:t[^>]*>(.*?)</w:t>")


def _extract_first_tag(fragment: str, tag: str) -> str:
    pattern = re.compile(rf"<{tag}[^>]*>(.*?)</{tag}>", re.DOTALL)
    m = pattern.search(fragment)
    if not m:
        return ""
    return _extract_text(m.group(1))


def _extract_text(xml_fragment: str) -> str:
    return "".join(TEXT_PATTERN.findall(xml_fragment))


def omml_to_latex_from_xml(xml: str) -> str:
    work = xml

    # 1. fractions
    def repl_frac(m: re.Match) -> str:
        frag = m.group(1)
        num_block = NUM_PATTERN.search(frag)
        den_block = DEN_PATTERN.search(frag)
        num = _extract_text(num_block.group(1)) if num_block else ""
        den = _extract_text(den_block.group(1)) if den_block else ""
        return f"\\frac{{{num}}}{{{den}}}"

    work = FRACTION_PATTERN.sub(lambda m: repl_frac(m), work)

    # 2. sub/sup combos
    def _extract_tag_text(fragment: str, tag: str) -> str:
        return _extract_first_tag(fragment, tag)

    def repl_subsup(pattern: re.Pattern, kind: str, frag: str) -> str:
        base = _extract_first_tag(frag, "m:e") or ""
        sub = _extract_first_tag(frag, "m:sub") or ""
        sup = _extract_first_tag(frag, "m:sup") or ""
        if kind == "sup":
            return f"{base}^{{{sup}}}" if sup else base
        if kind == "sub":
            return f"{base}_{{{sub}}}" if sub else base
        return f"{base}_{{{sub}}}^{{{sup}}}" if (sub or sup) else base

    def subsup_replacer(regex: re.Pattern, kind: str, text: str) -> str:
        def _inner(m: re.Match):
            return repl_subsup(regex, kind, m.group(1))

        return regex.sub(lambda m: _inner(m), text)

    work = subsup_replacer(SUBSUP_PATTERN, "subsup", work)
    work = subsup_replacer(SUP_PATTERN, "sup", work)
    work = subsup_replacer(SUB_PATTERN, "sub", work)

    # 3. radicals
    def repl_rad(m: re.Match) -> str:
        frag = m.group(1)
        deg = _extract_first_tag(frag, "m:deg")
        expr = _extract_first_tag(frag, "m:e")
        if deg:
            return f"\\sqrt[{deg}]{{{expr}}}"
        return f"\\sqrt{{{expr}}}"

    work = RAD_PATTERN.sub(lambda m: repl_rad(m), work)

    # 4. simple n-ary (sum/integral) -- we only map the operator symbol and body
    def repl_nary(m: re.Match) -> str:
        frag = m.group(1)
        # operator character appears as <m:chr m:val="∑" /> OR sometimes nested; just search
        chr_m = re.search(r'<m:chr[^>]*?m:val="([^"]+)"', frag)
        body = _extract_first_tag(frag, "m:e")
        op = chr_m.group(1) if chr_m else "∑"
        mapping = {"∑": "\\sum", "∫": "\\int", "∏": "\\prod"}
        op_latex = mapping.get(op, op)
        return f"{op_latex}({body})"

    work = NARY_PATTERN.sub(lambda m: repl_nary(m), work)

    # 5. Fallback: extract all text (w:t) for anything remaining of interest
    txt = _extract_text(work)
    return txt if txt.strip() else "/*math*/"


def has_math_xml(xml: str) -> bool:
    return "<m:oMath" in xml or "<m:oMathPara" in xml


def paragraph_with_math(p) -> str:
    xml = p._p.xml  # string
    text_plain = p.text or ""
    if not has_math_xml(xml):
        return text_plain.strip()

    # 简单将 math 块替换
    # 提取每个 <m:oMath...>...</m:oMath> / <m:oMathPara ...> ... </m:oMathPara>
    def repl(match: re.Match) -> str:
        return "$ " + omml_to_latex_from_xml(match.group(0)) + " $"

    pattern_math = re.compile(
        r"<m:oMath[^>]*>.*?</m:oMath>|<m:oMathPara[^>]*>.*?</m:oMathPara>", re.DOTALL
    )
    replaced = pattern_math.sub(repl, xml)
    # 再抽取所有 w:t 形成线性文本（其中 math 已替换为 LaTeX）
    # 临时占位符用特殊标记避免被剥离
    placeholder_map = {}
    placeholder_idx = 0

    def keep_math(m: re.Match):
        nonlocal placeholder_idx
        key = f"__MATH_{placeholder_idx}__"
        placeholder_idx += 1
        placeholder_map[key] = m.group(0)
        return key

    # 将 $ ... $ 片段临时占位
    replaced2 = re.sub(r"\$[^$]+\$", keep_math, replaced)
    texts = TEXT_PATTERN.findall(replaced2)
    merged = " ".join(t.strip() for t in texts if t.strip())
    # 还原占位符
    for k, v in placeholder_map.items():
        merged = merged.replace(k, v)
    # 保证 math 片段与文字间留空
    merged = re.sub(r"\s+", " ", merged).strip()
    return merged


def extract_paragraphs(doc: Document) -> Iterable[Dict]:
    for p in doc.paragraphs:
        if not (p.text or has_math_xml(p._p.xml)):
            continue
        yield {
            "type": "paragraph",
            "text": paragraph_with_math(p),
            "style": p.style.name if p.style else "",
        }
    for tbl_idx, tbl in enumerate(doc.tables):
        rows = []
        for r in tbl.rows:
            row_text = [cell.text.strip() for cell in r.cells]
            rows.append(row_text)
        yield {"type": "table", "rows": rows, "index": tbl_idx}


def load_docx(path: Path) -> List[Dict]:
    doc = Document(str(path))
    elements = list(extract_paragraphs(doc))
    # Attach source path
    for i, e in enumerate(elements):
        e["source"] = str(path)
        e["order"] = i
    return elements


PARSER_VERSION = "2025-09-01a"


def file_hash(path: Path, block_size: int = 65536) -> str:
    h = hashlib.sha256()
    try:
        with open(path, "rb") as f:
            for chunk in iter(lambda: f.read(block_size), b""):
                h.update(chunk)
        return h.hexdigest()[:16]
    except Exception:
        return ""


def iter_supported(root: Path) -> Iterable[Path]:
    """Yield supported files (docx + pdf)."""
    if not root.exists():
        return []
    for p in root.rglob("*"):
        if p.is_dir():
            continue
        if p.name.startswith("~$"):
            continue
        ext = p.suffix.lower()
        if ext == ".docx" or (ext == ".pdf" and HAVE_PDF):
            yield p


def list_docx_paths(root: Path) -> List[Path]:  # backward compatibility
    return [p for p in iter_supported(root) if p.suffix.lower() == ".docx"]


def list_supported_paths(root: Path) -> List[Path]:
    return list(iter_supported(root))


def _load_pdf(path: Path, max_mb: int = 25, low_ratio: float = 0.02) -> List[Dict]:
    if not HAVE_PDF:
        return [{"type": "error", "source": str(path), "error": "pdfminer 未安装"}]
    size_bytes = path.stat().st_size if path.exists() else 0
    if size_bytes > max_mb * 1024 * 1024:
        return [
            {
                "type": "error",
                "source": str(path),
                "error": f"pdf_oversize>{max_mb}MB",
                "size_bytes": size_bytes,
            }
        ]
    output = StringIO()
    try:
        # 抑制 pdfminer 在部分异常颜色指令上产生的大量噪声日志：
        # 典型信息: "Cannot set gray non-stroke color because /'P90' is an invalid float value"
        @contextmanager
        def _suppress_pdfminer_noise():
            target_logger_names = [
                "pdfminer",
                "pdfminer.pdfinterp",
                "pdfminer.converter",
                "pdfminer.layout",
            ]
            filters_added = []

            class _PdfNoiseFilter(logging.Filter):
                def filter(self, record: logging.LogRecord) -> bool:  # noqa: D401
                    msg = record.getMessage()
                    if "Cannot set gray non-stroke color" in msg:
                        return False
                    if "is an invalid float value" in msg and "gray non-stroke" in msg:
                        return False
                    return True

            old_levels = {}
            for name in target_logger_names:
                lg = logging.getLogger(name)
                old_levels[name] = lg.level
                # 统一降级到 ERROR，添加过滤器
                lg.setLevel(logging.ERROR)
                flt = _PdfNoiseFilter()
                lg.addFilter(flt)
                filters_added.append((lg, flt))
            try:
                yield
            finally:
                for lg, flt in filters_added:
                    try:
                        lg.removeFilter(flt)
                    except Exception:  # pragma: no cover
                        pass
                for name, lvl in old_levels.items():
                    logging.getLogger(name).setLevel(lvl)

        with open(path, "rb") as f, _suppress_pdfminer_noise():
            extract_text_to_fp(f, output, laparams=None, output_type="text", codec="utf-8")
        text = output.getvalue()
        text_bytes = len(text.encode("utf-8")) if text else 0
        ratio = (text_bytes / size_bytes) if size_bytes else 0
        if ratio < low_ratio:
            # 可能是扫描版或文本极少
            return [
                {
                    "type": "error",
                    "source": str(path),
                    "error": "pdf_low_text_ratio",
                    "size_bytes": size_bytes,
                    "text_bytes": text_bytes,
                    "ratio": ratio,
                }
            ]
    except Exception as e:  # pragma: no cover
        return [{"type": "error", "source": str(path), "error": str(e)}]
    parts = []
    # 简单按空行拆分段落
    for i, block in enumerate(re.split(r"\n\s*\n", text)):
        blk = block.strip()
        if not blk:
            continue
        parts.append(
            {"type": "paragraph", "text": blk, "source": str(path), "order": i, "origin": "pdf"}
        )
    return parts


def ingest_to_raw(root: str, max_pdf_mb: int = 25, low_pdf_text_ratio: float = 0.02) -> List[Dict]:
    # Resolve to absolute path relative to project root if necessary
    root_path = Path(root)
    if not root_path.is_absolute():
        # project root assumed as two levels up from this file (src/ingestion)
        proj_root = Path(__file__).resolve().parents[2]
        candidate = proj_root / root
        if candidate.exists():
            root_path = candidate
    all_items: List[Dict] = []
    for f in iter_supported(root_path):
        if f.suffix.lower() == ".docx":
            try:
                elems = load_docx(f)
                fhash = file_hash(f)
                mtime = int(f.stat().st_mtime)
                for e in elems:
                    e["file_hash"] = fhash
                    e["file_mtime"] = mtime
                    e["parser_version"] = PARSER_VERSION
                all_items.extend(elems)
            except Exception as e:
                all_items.append({"type": "error", "source": str(f), "error": str(e)})
        elif f.suffix.lower() == ".pdf":
            pdf_items = _load_pdf(f, max_mb=max_pdf_mb, low_ratio=low_pdf_text_ratio)
            fhash = file_hash(f)
            mtime = int(f.stat().st_mtime)
            for e in pdf_items:
                e["file_hash"] = fhash
                e["file_mtime"] = mtime
                e["parser_version"] = PARSER_VERSION
            all_items.extend(pdf_items)
    return all_items


def ingest_files(
    files: List[Path], max_pdf_mb: int = 25, low_pdf_text_ratio: float = 0.02
) -> List[Dict]:
    all_items: List[Dict] = []
    for f in files:
        if f.suffix.lower() == ".docx":
            try:
                elems = load_docx(f)
                fhash = file_hash(f)
                mtime = int(f.stat().st_mtime)
                for e in elems:
                    e["file_hash"] = fhash
                    e["file_mtime"] = mtime
                    e["parser_version"] = PARSER_VERSION
                all_items.extend(elems)
            except Exception as e:
                all_items.append({"type": "error", "source": str(f), "error": str(e)})
        elif f.suffix.lower() == ".pdf":
            pdf_items = _load_pdf(f, max_mb=max_pdf_mb, low_ratio=low_pdf_text_ratio)
            fhash = file_hash(f)
            mtime = int(f.stat().st_mtime)
            for e in pdf_items:
                e["file_hash"] = fhash
                e["file_mtime"] = mtime
                e["parser_version"] = PARSER_VERSION
            all_items.extend(pdf_items)
    return all_items


if __name__ == "__main__":
    import json
    import sys

    root = sys.argv[1] if len(sys.argv) > 1 else "2025国赛创新型算法+源代码汇总！"
    data = ingest_to_raw(root)
    print(json.dumps(data[:5], ensure_ascii=False, indent=2))
